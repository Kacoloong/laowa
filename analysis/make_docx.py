first_result = ((
    (0, 5, 6, 7),
    (6, 2, 0, 0),
    (1, 7, 4, 6),
    (4, 3, 3, 2),
), ((4, 25), (10, 85), (13, 111), (14, 120), (15, 133), (20, 171)))

sboxes = ('''
70 2c b3 c0 e4 57 ea ae 23 6b 45 a5 ed 4f 1d 92
86 af 7c 1f 3e dc 5e b a6 39 d5 5d d9 5a 51 6c
8b 9a fb b0 74 2b f0 84 df cb 34 76 6d a9 d1 4
14 3a de 11 32 9c 53 f2 fe cf c3 7a 24 e8 60 69
aa a0 a1 62 54 1e e0 64 10 0 a3 75 8a e6 9 dd
87 83 cd 90 73 f6 9d bf 52 d8 c8 c6 81 6f 13 63
e9 a7 9f bc 29 f9 2f b4 78 6 e7 71 d4 ab 88 8d
72 b9 f8 ac 36 2a 3c f1 40 d3 bb 43 15 ad 77 80
82 ec 27 e5 85 35 c 41 ef 93 19 21 e 4e 65 bd
b8 8f eb ce 30 5f c5 1a e1 ca 47 3d 1 d6 56 4d
d 66 cc 2d 12 20 b1 99 4c c2 7e 5 b7 31 17 d7
58 61 1b 1c f 16 18 22 44 b2 b5 91 8 a8 fc 50
d0 7d 89 97 5b 95 ff d2 c4 48 f7 db 3 da 3f 94
5c 2 4a 33 67 f3 7f e2 9b 26 37 3b 96 4b be 2e
79 8c 6e 8e f5 b6 fd 59 98 6a 46 ba 25 42 a2 fa
7 55 ee a 49 68 38 a4 28 7b c9 c1 e3 f4 c7 9e
''', '''
70 2c b3 c0 e4 57 ea ae 23 6b 45 a5 ed 4f 1d 92
86 af 7c 1f 3e dc 5e b a6 39 d5 5d d9 5a 51 6c
8b 9a fb b0 74 2b f0 84 df cb 34 76 6d a9 d1 4
14 3a de 11 32 9c 53 f2 fe cf c3 7a 24 e8 60 69
aa a0 a1 62 54 1e e0 64 10 0 a3 75 8a e6 9 dd
87 83 cd 90 73 f6 9d bf 52 d8 c8 c6 81 6f 13 63
e9 a7 9f bc 29 f9 2f b4 78 6 e7 71 d4 ab 88 8d
72 b9 f8 ac 36 2a 3c f1 40 d3 bb 43 15 ad 77 80
82 ec 27 e5 85 35 c 41 ef 93 19 21 e 4e 65 bd
b8 8f eb ce 30 5f c5 1a e1 ca 47 3d 1 d6 56 4d
d 66 cc 2d 12 20 b1 99 4c c2 7e 5 b7 31 17 d7
58 61 1b 1c f 16 18 22 44 b2 b5 91 8 a8 fc 50
d0 7d 89 97 5b 95 ff d2 c4 48 f7 db 3 da 3f 94
5c 2 4a 33 67 f3 7f e2 9b 26 37 3b 96 4b be 2e
79 8c 6e 8e f5 b6 fd 59 98 6a 46 ba 25 42 a2 fa
7 55 ee a 49 68 38 a4 28 7b c9 c1 e3 f4 c7 9e
''')

mdses = ((
    (1, 1, 140, 141),
    (1, 140, 141, 1),
    (140, 141, 1, 1),
    (141, 1, 1, 140),
),)

column_moves = (
    (2, 0, 1, 3),
)


def nid_tuple_to_str(data):
    return '\n'.join(' '.join(map(str, row)) for row in data)

def append_a_lane_transform_result(table, item, sbox, column_move, mds):
    r = table.add_row()
    cell0 = r.cells[0]
    cell1 = r.cells[1]
    cell2 = r.cells[2]
    cell3 = r.cells[3]
    cell0.text = sbox
    cell1.text = str(column_move)
    cell2.text = nid_tuple_to_str(
        item[0]
    )
    cell3.text = nid_tuple_to_str(mds)
    for round, number in item[1]:
        cells = r.cells
        cell0.merge(cells[0])
        cell1.merge(cells[1])
        cell2.merge(cells[2])
        cell3.merge(cells[3])
        cells[4].text = str(round)
        cells[5].text = str(number)
        r = table.add_row()

    # Access the table's XML element
    tbl = table._element

    # Access the last row's XML element
    last_row = table.rows[-1]._element

    tbl.remove(last_row)


from docx import Document
import re
import logging

logging.basicConfig(level = logging.DEBUG)

doc = Document()

doc.add_heading('胡泓震的活跃S盒结果')
for i in range(len(sboxes)):
    p = doc.add_paragraph('S')
    r = p.add_run(f'{i}')
    r.font.subscript = True
    p.add_run(f':\n{str(sboxes[i])}')

table = doc.add_table(1,6)
table.style = 'Table Grid'
table_head = table.rows[0].cells
table_head[0].text = '选用的S盒'
table_head[1].text = '列位移参数'
table_head[2].text = 'Lane位移参数'
table_head[3].text = 'Mds'
table_head[4].text = 'r'
table_head[5].text = '活跃S盒个数（合格线：171）'

results = dict()
round_regex = re.compile(r'Round: (\d+) lane_transform:')
number_regex = re.compile(r'no: (\d+).0 time: \d+')
with open('./active_sbox_results', 'r') as f:
    while (l := f.readline()) != '':
        logging.debug(l)
        round = round_regex.match(l)
        if round is None:
            raise ValueError()
        round = int(round.group(1))
        l = f.readline().strip()
        lane = eval(l.strip())
        l = f.readline().strip()
        logging.debug(l)
        number = number_regex.match(l)
        if number is None:
            raise ValueError()
        number = int(number.group(1))
        t = results.get(lane)
        if t:
            t.append((round, number))
        else:
            results[lane] = [(round, number)]

append_a_lane_transform_result(table, first_result, 'S1', column_moves[0], mdses[0])
for item in results.items():
    append_a_lane_transform_result(table, item, 'S1', column_moves[0], mdses[0])

doc.save('../活跃S盒结果(胡泓震).docx')
